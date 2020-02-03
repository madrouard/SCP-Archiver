from bs4 import BeautifulSoup
import requests
from docx import Document
import os

if not os.path.isdir('./scpArchive'):
    os.mkdir('./scpArchive')


for i in range(2,10):
    scp = requests.get("http://www.scp-wiki.net/scp-00" + str(i))
    soup = BeautifulSoup(scp.text, 'html.parser')
    main_text = soup.select('p')
    file = Document()
    for j in range(8, len(main_text)-1):
        try:
            file.add_paragraph(main_text[j].getText())
        except ValueError:
            file.add_paragraph('Line ' + str(j) + " had an error")
    file.save('./scpArchive/scp-00' + str(i) + ".docx")

for i in range(10,100):
    scp = requests.get("http://www.scp-wiki.net/scp-0" + str(i))
    soup = BeautifulSoup(scp.text, 'html.parser')
    main_text = soup.select('p')
    file = Document()
    for j in range(8, len(main_text)-1):
        try:
            file.add_paragraph(main_text[j].getText())
        except ValueError:
            file.add_paragraph('Line ' + str(j) + " had an error")
    file.save('./scpArchive/scp-0' + str(i) + ".docx")

for i in range(100,5000):
    scp = requests.get("http://www.scp-wiki.net/scp-" + str(i))
    soup = BeautifulSoup(scp.text, 'html.parser')
    main_text = soup.select('p')
    file = Document()
    for j in range(8, len(main_text)-1):
        try:
            file.add_paragraph(main_text[j].getText())
        except ValueError:
            file.add_paragraph('Line ' + str(j) + " had an error")
    file.save('./scpArchive/scp-' + str(i) + ".docx")
